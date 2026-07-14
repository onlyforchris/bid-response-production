# -*- coding: utf-8 -*-
"""Convert a Markdown technical bid draft into a formatted DOCX.

This is intentionally small: the agent writes the bid content, this script
applies the repeatable Word formatting baseline.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


FONT = "宋体"


def set_run_font(run, size=12, bold=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def shade(cell, fill="D9D9D9"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.strip())
    set_run_font(run, 12, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_para(doc, text, size=12, bold=False, align=None, indent=True):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(5)
    if indent:
        pf.first_line_indent = Pt(21)
    run = p.add_run(text.strip())
    set_run_font(run, size, bold)
    return p


def setup_styles(doc):
    specs = {
        "Normal": (12, False),
        "Heading 1": (16, True),
        "Heading 2": (14, True),
        "Heading 3": (12, True),
    }
    for name, (size, bold) in specs.items():
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4 if level == 2 else 3)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.strip())
    set_run_font(run, 16 if level == 1 else 14 if level == 2 else 12, True)
    return p


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        parts = [p.strip() for p in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", p) for p in parts):
            rows.append(parts)
        i += 1
    width = max((len(r) for r in rows), default=0)
    rows = [r + [""] * (width - len(r)) for r in rows]
    return rows, i


def add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            set_cell_text(cell, val, bold=(ri == 0))
            if ri == 0:
                shade(cell)
    doc.add_paragraph()


def add_cover(doc, title, bidder, date_text):
    for _ in range(5):
        doc.add_paragraph()
    add_para(doc, title, 22, True, WD_ALIGN_PARAGRAPH.CENTER, False)
    add_para(doc, "技术标书", 26, True, WD_ALIGN_PARAGRAPH.CENTER, False)
    for _ in range(3):
        doc.add_paragraph()
    add_para(doc, f"投标人：{bidder}", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
    add_para(doc, f"日期：{date_text}", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
    doc.add_page_break()


def convert(md_path: Path, out_path: Path, title: str, bidder: str, date_text: str):
    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    add_cover(doc, title, bidder, date_text)

    lines = md_path.read_text(encoding="utf-8-sig").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            add_heading(doc, m.group(2), min(len(m.group(1)), 3))
        elif re.match(r"^[-*]\s+", line):
            add_para(doc, re.sub(r"^[-*]\s+", "", line), indent=True)
        else:
            add_para(doc, line)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("markdown")
    ap.add_argument("output")
    ap.add_argument("--title", default="技术标书")
    ap.add_argument("--bidder", default="[投标人名称]")
    ap.add_argument("--date", default="")
    args = ap.parse_args()
    convert(Path(args.markdown), Path(args.output), args.title, args.bidder, args.date)


if __name__ == "__main__":
    main()
