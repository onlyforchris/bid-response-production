# -*- coding: utf-8 -*-
"""Quick structural review for a Chinese technical bid DOCX."""

from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


RISK_TERMS = [
    "[投标人名称]",
    "[客户名称]",
    "[项目名称]",
    "TODO",
    "待补",
    "待完善",
    "演示",
    "测试数据",
    "localhost",
    "127.0.0.1",
    "POC",
    "100%",
    "确保",
    "保证",
    "无误",
    "零风险",
]


def app_page_count(path: Path):
    with zipfile.ZipFile(path) as z:
        if "docProps/app.xml" not in z.namelist():
            return None
        app = z.read("docProps/app.xml").decode("utf-8", errors="ignore")
    m = re.search(r"<Pages>(\d+)</Pages>", app)
    return int(m.group(1)) if m else None


def font_name(run):
    rpr = run._element.rPr
    if rpr is not None and rpr.rFonts is not None:
        return rpr.rFonts.get(qn("w:eastAsia")) or run.font.name
    return run.font.name


def review(path: Path):
    doc = Document(str(path))
    para_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    table_text = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    table_text.append(cell.text.strip())
    text = "\n".join(para_text + table_text)

    sizes = Counter()
    fonts = Counter()
    line_spacing = Counter()
    heading_styles = Counter()
    for para in doc.paragraphs:
        run = next((r for r in para.runs if r.text.strip()), None)
        if not run:
            continue
        if para.style.name.startswith("Heading"):
            heading_styles[para.style.name] += 1
        fonts[font_name(run) or "unknown"] += 1
        sizes[str(run.font.size.pt if run.font.size else "inherited")] += 1
        line_spacing[str(para.paragraph_format.line_spacing or "inherited")] += 1

    risks = {term: text.count(term) for term in RISK_TERMS if text.count(term)}
    tables = [
        {
            "index": i,
            "rows": len(t.rows),
            "cols": len(t.columns),
            "first_row": [c.text.strip()[:80] for c in t.rows[0].cells] if t.rows else [],
        }
        for i, t in enumerate(doc.tables)
    ]

    module_lengths = []
    for idx, para in enumerate(para_text):
        if not re.match(r"^3\.\d+\s+", para):
            continue
        body = []
        for next_para in para_text[idx + 1 :]:
            if re.match(r"^[3-6]\.\d+\s+", next_para) or next_para.startswith("第四章"):
                break
            body.append(next_para)
        module_lengths.append(
            {"title": para, "body_chars": sum(len(x) for x in body), "paragraphs": len(body)}
        )

    warnings = []
    if risks:
        warnings.append("risk terms remain")
    if module_lengths:
        avg = sum(x["body_chars"] for x in module_lengths) // len(module_lengths)
        if avg < 700:
            warnings.append("functional modules are too thin for a formal technical bid")
        if any(x["body_chars"] < 300 for x in module_lengths):
            warnings.append("some functional modules are below the minimum expansion gate")
    pages = app_page_count(path)
    if pages == 1 and len(text) > 5000:
        warnings.append("stored page count is stale; open/update fields or render to confirm real pages")

    return {
        "file": str(path),
        "app_pages": pages,
        "paragraphs": len(doc.paragraphs),
        "nonempty_paragraphs": len(para_text),
        "tables": len(doc.tables),
        "text_chars": len(text),
        "warnings": warnings,
        "risk_terms": risks,
        "module_summary": {
            "count": len(module_lengths),
            "average_body_chars": (
                sum(x["body_chars"] for x in module_lengths) // len(module_lengths)
                if module_lengths
                else 0
            ),
            "min_body_chars": min((x["body_chars"] for x in module_lengths), default=0),
            "thin_items": [x for x in module_lengths if x["body_chars"] < 300],
            "items": module_lengths,
        },
        "top_fonts": fonts.most_common(8),
        "top_font_sizes": sizes.most_common(8),
        "top_line_spacing": line_spacing.most_common(8),
        "heading_styles": heading_styles.most_common(),
        "table_summary": tables,
        "first_paragraphs": para_text[:20],
    }


def main():
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    ap = argparse.ArgumentParser()
    ap.add_argument("docx")
    args = ap.parse_args()
    print(json.dumps(review(Path(args.docx)), ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
